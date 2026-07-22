"""
Universal file loader — dispatches by extension.
Returns LangChain Document objects with a consistent metadata contract.
"""

from __future__ import annotations

import base64
from pathlib import Path
from typing import Callable, Dict, List

import pandas as pd
from docx import Document as DocxDocument
from docx.table import Table
from langchain_core.documents import Document
from PIL import Image

from etl.pdf_loader import load_pdf


def _df_to_markdown(df: pd.DataFrame) -> str:
    headers = [str(col) for col in df.columns]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in df.itertuples(index=False):
        cells = [str(value).replace("|", "\\|") for value in row]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _docx_table_to_markdown(table: Table) -> str:
    rows: List[str] = []
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if row_idx == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def _load_pdf(file_path: str) -> List[Document]:
    docs = load_pdf(file_path)
    path = Path(file_path)
    return [
        Document(
            page_content=doc.page_content,
            metadata={
                "source": doc.metadata.get("source", path.name),
                "page": doc.metadata.get("page", 1),
                "source_type": "pdf",
                "content_type": "text",
            },
        )
        for doc in docs
    ]


def _load_docx(file_path: str) -> List[Document]:
    path = Path(file_path)
    docx = DocxDocument(str(path))
    documents: List[Document] = []

    for idx, paragraph in enumerate(docx.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "page": idx,
                    "source_type": "docx",
                    "content_type": "text",
                },
            )
        )

    for idx, table in enumerate(docx.tables, start=1):
        markdown = _docx_table_to_markdown(table).strip()
        if not markdown:
            continue
        documents.append(
            Document(
                page_content=markdown,
                metadata={
                    "source": path.name,
                    "page": idx,
                    "source_type": "docx",
                    "content_type": "table",
                },
            )
        )

    if not documents:
        raise ValueError(f"No extractable content found in {path.name}")

    return documents


def _load_excel(file_path: str) -> List[Document]:
    path = Path(file_path)
    sheets = pd.read_excel(path, sheet_name=None)
    documents: List[Document] = []

    for idx, (_, df) in enumerate(sheets.items(), start=1):
        if df.empty:
            continue
        documents.append(
            Document(
                page_content=_df_to_markdown(df),
                metadata={
                    "source": path.name,
                    "page": idx,
                    "source_type": "excel",
                    "content_type": "table",
                },
            )
        )

    if not documents:
        raise ValueError(f"No extractable content found in {path.name}")

    return documents


def _load_csv(file_path: str) -> List[Document]:
    path = Path(file_path)
    df = pd.read_csv(path)
    if df.empty:
        raise ValueError(f"No extractable content found in {path.name}")

    return [
        Document(
            page_content=_df_to_markdown(df),
            metadata={
                "source": path.name,
                "page": 1,
                "source_type": "excel",
                "content_type": "table",
            },
        )
    ]


def _load_image(file_path: str) -> List[Document]:
    path = Path(file_path)
    with Image.open(path) as img:
        img.verify()

    with Image.open(path) as img:
        img.load()

    image_bytes = path.read_bytes()

    return [
        Document(
            page_content=base64.b64encode(image_bytes).decode("ascii"),
            metadata={
                "source": path.name,
                "page": 1,
                "source_type": "image",
                "content_type": "image",
            },
        )
    ]


def _load_text(file_path: str) -> List[Document]:
    path = Path(file_path)
    content = path.read_text(encoding="utf-8")
    if not content.strip():
        raise ValueError(f"No extractable content found in {path.name}")

    return [
        Document(
            page_content=content,
            metadata={
                "source": path.name,
                "page": 1,
                "source_type": "text",
                "content_type": "text",
            },
        )
    ]


_LOADERS: Dict[str, Callable[[str], List[Document]]] = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".xlsx": _load_excel,
    ".xls": _load_excel,
    ".csv": _load_csv,
    ".jpg": _load_image,
    ".jpeg": _load_image,
    ".png": _load_image,
    ".webp": _load_image,
    ".txt": _load_text,
    ".md": _load_text,
}


def load_file(file_path: str) -> List[Document]:
    """
    Load a file and return Documents based on its extension.

    Args:
        file_path: Absolute or relative path to the file.

    Returns:
        List of Document objects with metadata keys:
        source, page, source_type, content_type.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        supported = ", ".join(sorted(_LOADERS))
        raise ValueError(
            f"Unsupported file type '{path.suffix}' for {path.name}. "
            f"Supported extensions: {supported}"
        )

    return loader(str(path))
